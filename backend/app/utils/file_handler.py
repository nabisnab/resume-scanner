"""
File handling utilities for parsing resumes in multiple formats.

This module provides utilities to extract text from various file formats:
- PDF files
- DOCX (Word) files
- TXT (Plain text) files
"""

import os
import logging
from pathlib import Path
from typing import Optional, Tuple
from io import BytesIO

logger = logging.getLogger(__name__)


class FileHandlerError(Exception):
    """Custom exception for file handling errors."""
    pass


class FileHandler:
    """Utility class for handling and parsing multiple file formats."""
    
    # Supported file extensions
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.txt'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
    
    @staticmethod
    def validate_file(file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Validate if a file exists and has a supported format.
        
        Args:
            file_path: Path to the file to validate
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            path = Path(file_path)
            
            # Check if file exists
            if not path.exists():
                return False, f"File not found: {file_path}"
            
            # Check if it's a file
            if not path.is_file():
                return False, f"Path is not a file: {file_path}"
            
            # Check file extension
            if path.suffix.lower() not in FileHandler.SUPPORTED_FORMATS:
                return False, f"Unsupported file format: {path.suffix}. Supported formats: {', '.join(FileHandler.SUPPORTED_FORMATS)}"
            
            # Check file size
            file_size = path.stat().st_size
            if file_size == 0:
                return False, "File is empty"
            
            if file_size > FileHandler.MAX_FILE_SIZE:
                return False, f"File size exceeds maximum limit of {FileHandler.MAX_FILE_SIZE / (1024 * 1024)} MB"
            
            return True, None
            
        except Exception as e:
            return False, f"Error validating file: {str(e)}"
    
    @staticmethod
    def parse_file(file_path: str) -> str:
        """
        Parse a file and extract text content.
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            Extracted text content
            
        Raises:
            FileHandlerError: If file cannot be parsed
        """
        is_valid, error_msg = FileHandler.validate_file(file_path)
        if not is_valid:
            raise FileHandlerError(error_msg)
        
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return FileHandler._parse_pdf(file_path)
            elif file_extension == '.docx':
                return FileHandler._parse_docx(file_path)
            elif file_extension == '.txt':
                return FileHandler._parse_txt(file_path)
            else:
                raise FileHandlerError(f"Unsupported format: {file_extension}")
                
        except FileHandlerError:
            raise
        except Exception as e:
            raise FileHandlerError(f"Error parsing {file_extension} file: {str(e)}")
    
    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
            
        Raises:
            FileHandlerError: If PDF cannot be parsed
        """
        try:
            import PyPDF2
        except ImportError:
            raise FileHandlerError("PyPDF2 library is required for PDF parsing. Install it with: pip install PyPDF2")
        
        text_content = []
        
        try:
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                if len(pdf_reader.pages) == 0:
                    raise FileHandlerError("PDF file has no pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from PDF page {page_num + 1}: {str(e)}")
                        continue
            
            if not text_content:
                raise FileHandlerError("No text could be extracted from the PDF file")
            
            return '\n'.join(text_content)
            
        except FileHandlerError:
            raise
        except Exception as e:
            raise FileHandlerError(f"Failed to read PDF file: {str(e)}")
    
    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """
        Extract text from a DOCX (Word) file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            FileHandlerError: If DOCX cannot be parsed
        """
        try:
            from docx import Document
        except ImportError:
            raise FileHandlerError("python-docx library is required for DOCX parsing. Install it with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if not text_content:
                raise FileHandlerError("No text could be extracted from the DOCX file")
            
            return '\n'.join(text_content)
            
        except FileHandlerError:
            raise
        except Exception as e:
            raise FileHandlerError(f"Failed to parse DOCX file: {str(e)}")
    
    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """
        Extract text from a TXT file.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            File content
            
        Raises:
            FileHandlerError: If TXT file cannot be read
        """
        try:
            # Try UTF-8 encoding first
            with open(file_path, 'r', encoding='utf-8') as txt_file:
                content = txt_file.read()
            
            if not content.strip():
                raise FileHandlerError("TXT file is empty or contains only whitespace")
            
            return content
            
        except UnicodeDecodeError:
            # Fallback to latin-1 encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as txt_file:
                    content = txt_file.read()
                
                if not content.strip():
                    raise FileHandlerError("TXT file is empty or contains only whitespace")
                
                return content
                
            except Exception as e:
                raise FileHandlerError(f"Failed to read TXT file with any encoding: {str(e)}")
        
        except FileHandlerError:
            raise
        except Exception as e:
            raise FileHandlerError(f"Failed to read TXT file: {str(e)}")
    
    @staticmethod
    def parse_from_bytes(file_bytes: bytes, file_name: str) -> str:
        """
        Parse a file from bytes content.
        
        Args:
            file_bytes: Raw file bytes
            file_name: Name of the file (used to determine format)
            
        Returns:
            Extracted text content
            
        Raises:
            FileHandlerError: If file cannot be parsed
        """
        file_extension = Path(file_name).suffix.lower()
        
        if file_extension not in FileHandler.SUPPORTED_FORMATS:
            raise FileHandlerError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == '.pdf':
                return FileHandler._parse_pdf_from_bytes(file_bytes)
            elif file_extension == '.docx':
                return FileHandler._parse_docx_from_bytes(file_bytes)
            elif file_extension == '.txt':
                return FileHandler._parse_txt_from_bytes(file_bytes)
            else:
                raise FileHandlerError(f"Unsupported format: {file_extension}")
                
        except FileHandlerError:
            raise
        except Exception as e:
            raise FileHandlerError(f"Error parsing {file_extension} from bytes: {str(e)}")
    
    @staticmethod
    def _parse_pdf_from_bytes(file_bytes: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            import PyPDF2
        except ImportError:
            raise FileHandlerError("PyPDF2 library is required for PDF parsing. Install it with: pip install PyPDF2")
        
        text_content = []
        
        try:
            pdf_file = BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            if len(pdf_reader.pages) == 0:
                raise FileHandlerError("PDF file has no pages")
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page {page_num + 1}: {str(e)}")
                    continue
            
            if not text_content:
                raise FileHandlerError("No text could be extracted from the PDF file")
            
            return '\n'.join(text_content)
            
        except FileHandlerError:
            raise
        except Exception as e:
            raise FileHandlerError(f"Failed to parse PDF from bytes: {str(e)}")
    
    @staticmethod
    def _parse_docx_from_bytes(file_bytes: bytes) -> str:
        """Extract text from DOCX bytes."""
        try:
            from docx import Document
        except ImportError:
            raise FileHandlerError("python-docx library is required for DOCX parsing. Install it with: pip install python-docx")
        
        try:
            doc_file = BytesIO(file_bytes)
            doc = Document(doc_file)
            text_content = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if not text_content:
                raise FileHandlerError("No text could be extracted from the DOCX file")
            
            return '\n'.join(text_content)
            
        except FileHandlerError:
            raise
        except Exception as e:
            raise FileHandlerError(f"Failed to parse DOCX from bytes: {str(e)}")
    
    @staticmethod
    def _parse_txt_from_bytes(file_bytes: bytes) -> str:
        """Extract text from TXT bytes."""
        try:
            # Try UTF-8 encoding first
            content = file_bytes.decode('utf-8')
            
            if not content.strip():
                raise FileHandlerError("TXT file is empty or contains only whitespace")
            
            return content
            
        except UnicodeDecodeError:
            # Fallback to latin-1 encoding
            try:
                content = file_bytes.decode('latin-1')
                
                if not content.strip():
                    raise FileHandlerError("TXT file is empty or contains only whitespace")
                
                return content
                
            except Exception as e:
                raise FileHandlerError(f"Failed to decode TXT from bytes with any encoding: {str(e)}")
        
        except FileHandlerError:
            raise
        except Exception as e:
            raise FileHandlerError(f"Failed to parse TXT from bytes: {str(e)}")
